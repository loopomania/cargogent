#!/usr/bin/env python3
"""Convert Excel-format.md to Word document."""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt

def parse_markdown_table(lines):
    """Parse markdown table into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or not line.startswith("|"):
            continue
        cells = [c.strip().replace("**", "") for c in line.split("|")[1:-1]]
        if cells and not all(re.match(r'^[-]+$', c) for c in cells):
            rows.append(cells)
    return rows

def md_to_docx(md_path, docx_path):
    doc = Document()
    doc.add_heading("Excel Format Specification", 0)

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    i = 0
    table_lines = []

    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            i += 1
            continue
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("---"):
            i += 1
            continue
        elif line.strip().startswith("|") and "|" in line:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_markdown_table(table_lines)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = cell
                doc.add_paragraph()
            table_lines = []
            continue
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.strip()[2:].replace("**", ""))
        elif line.strip():
            doc.add_paragraph(line.strip().replace("**", ""))
        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")

if __name__ == "__main__":
    base = Path(__file__).parent.parent
    md_path = base / "Excel-format.md"
    docx_path = base / "Excel-format.docx"
    md_to_docx(md_path, docx_path)
